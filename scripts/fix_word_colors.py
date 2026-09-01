#!/usr/bin/env python3
"""Fix BiDi, patient-table row layout, header, and footer signature placement."""

from __future__ import annotations

import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SRC = "/opt/cursor/artifacts/word_converted_input.docx"
OUT = "/opt/cursor/artifacts/medical_certificate.docx"

BLUE = RGBColor(0x00, 0x00, 0x8B)
BLACK = RGBColor(0x00, 0x00, 0x00)
LRM = "\u200e"
RLM = "\u200f"


def has_drawing(run) -> bool:
    return bool(run._element.xpath(".//w:drawing"))


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        pPr.append(bidi)


def set_hebrew_lang(run):
    rPr = run._element.get_or_add_rPr()
    lang = rPr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rPr.append(lang)
    lang.set(qn("w:bidi"), "he-IL")
    lang.set(qn("w:rtl"), "he-IL")


def set_rtl_run(run):
    rPr = run._element.get_or_add_rPr()
    existing = rPr.find(qn("w:rtl"))
    if existing is not None:
        rPr.remove(existing)
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)
    set_hebrew_lang(run)


def set_ltr_run(run):
    rPr = run._element.get_or_add_rPr()
    existing = rPr.find(qn("w:rtl"))
    if existing is not None:
        rPr.remove(existing)
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "0")
    rPr.append(rtl)


def add_run(
    paragraph,
    text,
    *,
    blue=False,
    bold=False,
    underline=False,
    ltr=False,
    rtl=False,
    size_pt=None,
):
    r = paragraph.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.color.rgb = BLUE if blue else BLACK
    if size_pt:
        r.font.size = Pt(size_pt)
    if underline:
        r.font.underline = WD_UNDERLINE.SINGLE
    if ltr:
        set_ltr_run(r)
    elif rtl:
        set_rtl_run(r)
    return r


def add_rtl(paragraph, text, **kwargs):
    return add_run(paragraph, text, rtl=True, **kwargs)


def add_ltr(paragraph, text):
    add_run(paragraph, LRM + text, ltr=True)
    add_run(paragraph, RLM)


def clear_text_runs(paragraph):
    for run in list(paragraph.runs):
        if not has_drawing(run):
            run._element.getparent().remove(run._element)


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def no_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(borders)


def insert_header_table(doc):
    """Three-column header: date/referrer | Maccabi title | doctor name."""
    anchor = doc.paragraphs[0]._element
    ht = doc.add_table(rows=1, cols=3)
    no_border(ht)
    ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ht.autofit = False
    for cell in ht.rows[0].cells:
        cell.width = Cm(6)

    left = ht.rows[0].cells[2].paragraphs[0]
    set_rtl(left)
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_rtl(left, "תאריך: ", blue=True, bold=True)
    add_ltr(left, "31/08/2026")
    add_run(left, "\n")
    add_rtl(left, "מ.ר ", blue=True, bold=True)
    add_rtl(left, "גורם מפנה", blue=True, bold=True)

    center = ht.rows[0].cells[1].paragraphs[0]
    set_rtl(center)
    center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rtl(center, "מכבי שירותי בריאות", bold=True, size_pt=14)

    right = ht.rows[0].cells[0].paragraphs[0]
    set_rtl(right)
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_rtl(right, "דר' קגנוביץ גלינה", bold=True)

    ht._tbl.getparent().remove(ht._tbl)
    anchor.addprevious(ht._tbl)

    # Remove the mashed first paragraph (title/date/doctor were combined there).
    delete_paragraph(doc.paragraphs[0])


def fix_header_specialty(paragraph):
    set_rtl(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(4)
    clear_text_runs(paragraph)
    add_rtl(paragraph, "משפחה, פנימית וכללית")


def fix_contact(doc):
    phone_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "טלפון" in p.text or p.text.strip().startswith("טלפון"):
            phone_idx = i
            break
    if phone_idx is None:
        return
    p_phone = doc.paragraphs[phone_idx]
    p_addr = doc.paragraphs[phone_idx + 1]

    clear_text_runs(p_phone)
    set_rtl(p_phone)
    p_phone.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_phone.paragraph_format.space_after = Pt(2)
    add_rtl(p_phone, "טלפון: ", blue=True, bold=True)
    add_ltr(p_phone, "08-9257689")
    add_run(p_phone, "\n")
    add_rtl(p_phone, "פקס: ", blue=True, bold=True)
    add_ltr(p_phone, "08-9257689")

    clear_text_runs(p_addr)
    set_rtl(p_addr)
    p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_addr.paragraph_format.space_after = Pt(10)
    add_rtl(p_addr, "כתובת: ", blue=True, bold=True)
    add_rtl(p_addr, "שטרן יאיר 14, רמלה")


def cell_write(cell, lines: list[tuple[str, str, bool]]):
    """Each line: (label, value, value_is_ltr)."""
    p = cell.paragraphs[0]
    clear_text_runs(p)
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.15
    for i, (label, value, ltr_val) in enumerate(lines):
        if i:
            add_run(p, "\n")
        add_rtl(p, label, blue=True, bold=True)
        if value:
            add_run(p, " ")
            if ltr_val:
                add_ltr(p, value)
            else:
                add_rtl(p, value)


def patient_table(doc):
    for table in doc.tables:
        if len(table.columns) >= 4 and len(table.rows) >= 2:
            return table
    return doc.tables[-1]


def rebuild_patient_table(doc):
    old = patient_table(doc)
    drawing_elements = [
        r._element
        for row in old.rows
        for cell in row.cells
        for p in cell.paragraphs
        for r in p.runs
        if has_drawing(r)
    ]
    # Deduplicate drawing elements by id
    seen = set()
    unique_drawings = []
    for el in drawing_elements:
        key = id(el)
        if key not in seen:
            seen.add(key)
            unique_drawings.append(el)

    tbl_el = old._tbl
    parent = tbl_el.getparent()
    idx = list(parent).index(tbl_el)

    new = doc.add_table(rows=4, cols=4)
    parent.insert(idx, new._tbl)
    parent.remove(tbl_el)

    # Preserve blue border styling from Word conversion
    old_pr = tbl_el.find(qn("w:tblPr"))
    if old_pr is not None:
        new_pr = new._tbl.find(qn("w:tblPr"))
        if new_pr is None:
            new._tbl.insert(0, OxmlElement("w:tblPr"))
            new_pr = new._tbl.find(qn("w:tblPr"))
        for tag in ("w:tblBorders", "w:tblCellMar", "w:tblW"):
            el = old_pr.find(qn(tag))
            if el is not None:
                old_copy = new_pr.find(qn(tag))
                if old_copy is not None:
                    new_pr.remove(old_copy)
                new_pr.append(el)

    cell_write(new.rows[0].cells[3], [("פרטי הנבדק:", "", False)])
    cell_write(new.rows[1].cells[2], [("שם משפחה:", "גורפינקל", False)])
    cell_write(new.rows[1].cells[1], [("שם פרטי:", "רון", False)])
    cell_write(new.rows[1].cells[0], [("ת.ז.:", "215359910", True)])

    cell_write(
        new.rows[2].cells[2],
        [("ת.לידה:", "30/05/2005", True), ("מין:", "ז", False)],
    )
    cell_write(new.rows[2].cells[1], [("טלפון:", "08-9244787", True)])
    cell_write(new.rows[2].cells[0], [("טל.עבודה/נייד:", "0522460193", True)])

    cell_write(new.rows[3].cells[2], [("כתובת:", "הרב אפריאט 6, רמלה", False)])
    cell_write(new.rows[3].cells[0], [("מיקוד:", "7228212", True)])

    if unique_drawings:
        new.rows[1].cells[3].paragraphs[0]._p.append(unique_drawings[0])
    add_ltr(new.rows[3].cells[3].paragraphs[0], "0215359910")


def fix_title(p):
    clear_text_runs(p)
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    add_rtl(p, "אישור מחלה", bold=True, underline=True)


def set_doc_rtl(doc):
    """Default the document body to RTL (Hebrew)."""
    body = doc.element.body
    bodyPr = body.find(qn("w:bodyPr"))
    if bodyPr is None:
        bodyPr = OxmlElement("w:bodyPr")
        body.insert(0, bodyPr)
    if bodyPr.find(qn("w:bidi")) is None:
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        bodyPr.append(bidi)


def fix_statement(p):
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        if not has_drawing(r):
            r.font.color.rgb = BLUE
            r.bold = True
            set_rtl_run(r)
    p.paragraph_format.space_after = Pt(8)


def fix_dates(p):
    clear_text_runs(p)
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(20)
    add_rtl(p, "אינו/ה מסוגל/ת לעבוד מיום: ", blue=True, bold=True)
    add_ltr(p, "24/08/2026")
    add_run(p, "    ")
    add_rtl(p, "עד יום: ", blue=True, bold=True)
    add_ltr(p, "27/08/2026")
    add_run(p, "    ")
    add_rtl(p, 'סה"כ: ', blue=True, bold=True)
    add_rtl(p, "4 ימים.")


def fix_footer(doc):
    """Match PDF: date+line+label on RIGHT; signature line+label on LEFT."""
    footer_drawings = []
    to_remove = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        is_footer = txt in ("31/08/2026",) or "חתימה וחותמת" in txt or txt == "תאריך"
        is_empty_sig = not txt and any(has_drawing(r) for r in p.runs)
        if is_footer or is_empty_sig:
            for r in p.runs:
                if has_drawing(r):
                    footer_drawings.append(r._element)
            to_remove.append(p)

    for p in to_remove:
        delete_paragraph(p)

    ft = doc.add_table(rows=3, cols=2)
    no_border(ft)
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ft.autofit = False
    for row in ft.rows:
        row.cells[0].width = Cm(9)
        row.cells[1].width = Cm(9)

    # RTL col 0 = visual right (date side per PDF)
    date0 = ft.rows[0].cells[0].paragraphs[0]
    set_rtl(date0)
    date0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_ltr(date0, "31/08/2026")

    sig1 = ft.rows[1].cells[1].paragraphs[0]
    date1 = ft.rows[1].cells[0].paragraphs[0]
    set_rtl(sig1)
    set_rtl(date1)
    sig1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if footer_drawings:
        sig1._p.append(footer_drawings[0])
    else:
        add_run(sig1, "________________________")
    add_run(date1, "________________________")

    sig2 = ft.rows[2].cells[1].paragraphs[0]
    date2 = ft.rows[2].cells[0].paragraphs[0]
    set_rtl(sig2)
    set_rtl(date2)
    sig2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_rtl(sig2, "חתימה וחותמת הרופא", blue=True, bold=True)
    add_rtl(date2, "תאריך", blue=True, bold=True)

    date0.paragraph_format.space_before = Pt(24)


def remove_et4d(doc):
    for p in list(doc.paragraphs):
        if "ET4D" in p.text:
            delete_paragraph(p)


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    set_doc_rtl(doc)

    insert_header_table(doc)
    fix_header_specialty(doc.paragraphs[0])
    fix_contact(doc)
    rebuild_patient_table(doc)
    title_p = next(p for p in doc.paragraphs if "אישור מחלה" in p.text)
    fix_title(title_p)
    remove_et4d(doc)

    for p in doc.paragraphs:
        if "הנני לאשר" in p.text:
            fix_statement(p)
        if "מסוגל" in p.text:
            fix_dates(p)

    fix_footer(doc)
    doc.save(OUT)
    print("Saved", OUT)


if __name__ == "__main__":
    main()
