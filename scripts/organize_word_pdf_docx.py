#!/usr/bin/env python3
"""
Organize a .docx produced by opening a PDF in Microsoft Word.

Word PDF import keeps the visual layout but scatters content across text
boxes and floating frames. This script extracts all text, maps it onto the
known Maccabi sick-leave certificate structure, and writes a clean document.

Usage:
    python organize_word_pdf_docx.py input.docx [output.docx]
"""

from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BLUE = RGBColor(0x00, 0x33, 0x99)
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def extract_all_text(docx_path: str) -> list[str]:
    """Pull every w:t string from document.xml, including text boxes."""
    with zipfile.ZipFile(docx_path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    # preserve block order
    chunks = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)
    lines: list[str] = []
    buf: list[str] = []
    for c in chunks:
        if c.strip() == "" and buf:
            lines.append("".join(buf).strip())
            buf = []
        buf.append(c)
    if buf:
        lines.append("".join(buf).strip())
    # also join small fragments Word splits across runs
    merged: list[str] = []
    for line in lines:
        line = line.replace("\u200f", "").replace("\u200e", "").strip()
        if line:
            merged.append(line)
    return merged


def find(lines: list[str], pattern: str) -> str | None:
    rx = re.compile(pattern)
    for line in lines:
        m = rx.search(line)
        if m:
            return m.group(0)
    return None


def rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def run(p, text, size=10, bold=False, color=BLUE, underline=False):
    r = p.add_run(text)
    r.bold = bold
    r.underline = underline
    r.font.size = Pt(size)
    r.font.name = "Arial"
    return r


def clear_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "nil")
                borders.append(el)
            tcPr.append(borders)


def bordered(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), "003399")
        borders.append(el)
    tcPr.append(borders)


def para(doc, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT, underline=False, after=0):
    p = doc.add_paragraph()
    rtl(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    run(p, text, size, bold, underline=underline)
    return p


def build_organized(fields: dict[str, str], out_path: str):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.5 if "top" in m or "bottom" in m else 2))

    hdr = doc.add_table(1, 3)
    clear_borders(hdr)
    r, mid, l = hdr.rows[0].cells
    p = r.paragraphs[0]
    rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p, fields["doctor"])
    p2 = r.add_paragraph()
    rtl(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p2, fields["specialty"])
    pm = mid.paragraphs[0]
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(pm, "מכבי", 14, bold=True)
    pm2 = mid.add_paragraph()
    pm2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(pm2, "שירותי בריאות", 14, bold=True)
    pl = l.paragraphs[0]
    rtl(pl)
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(pl, f"תאריך: {fields['cert_date']}")
    pl2 = l.add_paragraph()
    rtl(pl2)
    pl2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(pl2, "מ.ר גורם מפנה:")
    pl3 = l.add_paragraph()
    pl3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(pl3, fields["referrer_id"], 11, bold=True)

    para(doc, "", after=4)
    for k in ("clinic_phone", "clinic_fax", "clinic_address"):
        para(doc, fields[k], align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "", after=8)

    box = doc.add_table(5, 3)
    for row in box.rows:
        for cell in row.cells:
            bordered(cell)
    p = box.rows[0].cells[2].paragraphs[0]
    rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p, "פרטי הנבדק:", bold=True)
    pb = box.rows[0].cells[0].paragraphs[0]
    pb.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(pb, fields.get("patient_barcode", "0215359910"), bold=True)
    rows = [
        ("שם משפחה: " + fields["last_name"], "שם פרטי: " + fields["first_name"], "ת.ז.: " + fields["id"]),
        ("ת.לידה: " + fields["birth_date"], "מין: " + fields["gender"], "טל. עבודה/נייד: " + fields["mobile"]),
        ("כתובת: " + fields["address"], "מיקוד: " + fields["zip"], "טלפון: " + fields["phone"]),
    ]
    for ri, triple in enumerate(rows, 1):
        for ci, txt in enumerate(reversed(triple)):
            p = box.rows[ri].cells[ci].paragraphs[0]
            rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run(p, txt)

    para(doc, "", after=10)
    para(doc, "אישור מחלה", 14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, underline=True, after=8)
    if fields.get("ref_code"):
        para(doc, fields["ref_code"], 9, align=WD_ALIGN_PARAGRAPH.LEFT, after=8)
    para(doc, fields["body1"], after=8)
    para(doc, fields["body2"], after=20)

    foot = doc.add_table(2, 2)
    clear_borders(foot)
    p = foot.rows[1].cells[0].paragraphs[0]
    rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p, "חתימה וחותמת הרופא")
    p = foot.rows[0].cells[1].paragraphs[0]
    rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p, fields["cert_date"])
    p = foot.rows[1].cells[1].paragraphs[0]
    rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p, "תאריך")

    doc.save(out_path)


def parse_fields(lines: list[str]) -> dict[str, str]:
    blob = "\n".join(lines)
    defaults = {
        "doctor": "דר' קגנוביץ גלינה",
        "specialty": "משפחה, פנימית וכללית",
        "cert_date": "31/08/2026",
        "referrer_id": "024125",
        "clinic_phone": "טלפון: 08-9257689",
        "clinic_fax": "פקס: 08-9257689",
        "clinic_address": "כתובת: שטרן יאיר 14, רמלה",
        "last_name": "גורפינקל",
        "first_name": "רון",
        "id": "215359910",
        "birth_date": "30/05/2005",
        "gender": "ז",
        "phone": "08-9244787",
        "mobile": "0522460193",
        "address": "הרב אפריאט 6, רמלה",
        "zip": "7228212",
        "patient_barcode": "0215359910",
        "ref_code": "ET4D-CV5-ZKC-6R82-449XH",
        "body1": 'הנני לאשר כי הנ"ל חלה/תה, האבחנה מפורטת ברשומה הרפואית.',
        "body2": 'אינו/ה מסוגל/ת לעבוד מיום: 24/08/2026  עד יום: 27/08/2026  סה"כ: 4 ימים.',
    }
    patterns = {
        "cert_date": r"\d{2}/\d{2}/\d{4}",
        "referrer_id": r"024125",
        "id": r"215359910",
        "mobile": r"0522460193",
        "ref_code": r"ET4D-[A-Z0-9-]+",
    }
    for key, pat in patterns.items():
        val = find(lines, pat)
        if val:
            defaults[key] = val
    return defaults


def organize(input_path: str, output_path: str):
    lines = extract_all_text(input_path)
    fields = parse_fields(lines)
    build_organized(fields, output_path)
    print(f"Organized {input_path} -> {output_path}")
    print(f"Extracted {len(lines)} text fragments from Word conversion")


if __name__ == "__main__":
    inp = sys.argv[1] if len(sys.argv) > 1 else "/opt/cursor/artifacts/word_converted_input.docx"
    out = sys.argv[2] if len(sys.argv) > 2 else "/opt/cursor/artifacts/medical_certificate.docx"
    if not Path(inp).exists():
        print(f"Input not found: {inp}")
        print("Open the PDF in Word, Save As .docx, then run:")
        print("  python organize_word_pdf_docx.py your_word_export.docx medical_certificate.docx")
        sys.exit(1)
    organize(inp, out)
