#!/usr/bin/env python3
"""Build an organized Word version of the Maccabi sick-leave certificate."""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = "/opt/cursor/artifacts/medical_certificate.docx"
BLUE = RGBColor(0x00, 0x33, 0x99)


def rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def run(p, text, size=10, bold=False, color=None, underline=False):
    r = p.add_run(text)
    r.bold = bold
    r.underline = underline
    r.font.size = Pt(size)
    r.font.name = "Arial"
    rf = r._element.get_or_add_rPr().get_or_add_rFonts()
    rf.set(qn("w:ascii"), "Arial")
    rf.set(qn("w:hAnsi"), "Arial")
    rf.set(qn("w:cs"), "Arial")
    if color:
        r.font.color.rgb = color
    return r


def para(doc, text="", size=10, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT, color=None, underline=False, space_after=0):
    p = doc.add_paragraph()
    rtl(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        run(p, text, size, bold, color, underline)
    return p


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val.get("val", "single"))
        el.set(qn("w:sz"), str(val.get("sz", 8)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), val.get("color", "003399"))
        borders.append(el)
    tcPr.append(borders)


def clear_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
            )


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)

    # ── Header row ──────────────────────────────────────────────
    hdr = doc.add_table(rows=1, cols=3)
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(hdr)
    r, mid, l = hdr.rows[0].cells

    # Right: doctor
    p = r.paragraphs[0]
    rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p, "דר' קגנוביץ גלינה", 10, color=BLUE)
    p2 = r.add_paragraph()
    rtl(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p2, "משפחה, פנימית וכללית", 10, color=BLUE)

    # Center: Maccabi
    pm = mid.paragraphs[0]
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(pm, "מכבי", 14, bold=True, color=BLUE)
    pm2 = mid.add_paragraph()
    pm2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(pm2, "שירותי בריאות", 14, bold=True, color=BLUE)

    # Left: date + referring entity
    pl = l.paragraphs[0]
    rtl(pl)
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(pl, "תאריך: 31/08/2026", 10, color=BLUE)
    pl2 = l.add_paragraph()
    rtl(pl2)
    pl2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(pl2, "מ.ר גורם מפנה:", 10, color=BLUE)
    pl3 = l.add_paragraph()
    pl3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(pl3, "024125", 11, bold=True)

    para(doc, space_after=4)

    # ── Clinic contact (centred) ────────────────────────────────
    for line in [
        "טלפון: 08-9257689",
        "פקס: 08-9257689",
        "כתובת: שטרן יאיר 14, רמלה",
    ]:
        para(doc, line, 10, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, space_after=2)

    para(doc, space_after=8)

    # ── Patient details box ───────────────────────────────────────
    box = doc.add_table(rows=5, cols=3)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in box.rows:
        for cell in row.cells:
            set_cell_border(cell)

    # Title row spanning feel
    title_cell = box.rows[0].cells[2]
    p = title_cell.paragraphs[0]
    rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p, "פרטי הנבדק:", 10, bold=True, color=BLUE)
    bc = box.rows[0].cells[0]
    pb = bc.paragraphs[0]
    pb.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(pb, "0215359910", 10, bold=True)

    patient = [
        ("שם משפחה: גורפינקל", "שם פרטי: רון", "ת.ז.: 215359910"),
        ("ת.לידה: 30/05/2005", "מין: ז", "טל. עבודה/נייד: 0522460193"),
        ("כתובת: הרב אפריאט 6, רמלה", "מיקוד: 7228212", "טלפון: 08-9244787"),
    ]
    for ri, (c2, c1, c0) in enumerate(patient, start=1):
        for ci, txt in enumerate([c0, c1, c2]):
            p = box.rows[ri].cells[ci].paragraphs[0]
            rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run(p, txt, 10, color=BLUE)

    para(doc, space_after=10)

    # ── Title ─────────────────────────────────────────────────────
    para(doc, "אישור מחלה", 14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, underline=True, space_after=8)

    # ── Reference code ────────────────────────────────────────────
    para(doc, "ET4D-CV5-ZKC-6R82-449XH", 9, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)

    # ── Body ──────────────────────────────────────────────────────
    para(
        doc,
        'הנני לאשר כי הנ"ל חלה/תה, האבחנה מפורטת ברשומה הרפואית.',
        10,
        space_after=8,
    )
    para(
        doc,
        'אינו/ה מסוגל/ת לעבוד מיום: 24/08/2026  עד יום: 27/08/2026  סה"כ: 4 ימים.',
        10,
        space_after=20,
    )

    # ── Footer: date + signature lines ───────────────────────────
    foot = doc.add_table(rows=2, cols=2)
    clear_table_borders(foot)

    sig = foot.rows[1].cells[0].paragraphs[0]
    rtl(sig)
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(sig, "חתימה וחותמת הרופא", 10, color=BLUE)

    dt = foot.rows[0].cells[1].paragraphs[0]
    rtl(dt)
    dt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(dt, "31/08/2026", 10, color=BLUE)
    dt2 = foot.rows[1].cells[1].paragraphs[0]
    rtl(dt2)
    dt2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(dt2, "תאריך", 10, color=BLUE)

    doc.save(OUT)
    print("Saved", OUT)


if __name__ == "__main__":
    main()
