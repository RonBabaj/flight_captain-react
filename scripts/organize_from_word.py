#!/usr/bin/env python3
"""Organize Word/Google-Docs PDF conversion to pixel-match the original layout."""

from __future__ import annotations

import os
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

INPUT = os.environ.get(
    "WORD_INPUT",
    "/opt/cursor/artifacts/word_converted_input.docx",
)
OUT = os.environ.get(
    "WORD_OUTPUT",
    "/opt/cursor/artifacts/medical_certificate.docx",
)
STAMP = "/opt/cursor/artifacts/image0.png"

BLUE = RGBColor(0x00, 0x33, 0x99)


def rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def txt(p, text, size=10, bold=False, color=BLUE, underline=False, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p.alignment = align
    rtl(p)
    r = p.add_run(text)
    r.bold = bold
    r.underline = underline
    r.font.size = Pt(size)
    r.font.name = "Arial"
    if color:
        r.font.color.rgb = color
    return r


def no_border(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            b = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "nil")
                b.append(el)
            tcPr.append(b)


def blue_border(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            b = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "8")
                el.set(qn("w:color"), "003399")
                b.append(el)
            tcPr.append(b)


def fix_phone(s: str) -> str:
    s = s.replace("\t", " ").strip()
    m = re.search(r"(\d{7,8})-(\d{2})", s)
    if m:
        return f"08-{m.group(1)}"
    m = re.search(r"08-?\d{7}", s.replace(" ", ""))
    if m:
        return m.group(0).replace("08", "08-", 1) if not m.group(0).startswith("08-") else m.group(0)
    return s


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.27)
    sec.bottom_margin = Cm(1.27)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)

    # ── Header ────────────────────────────────────────────────────
    hdr = doc.add_table(1, 3)
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_border(hdr)
    right, center, left = hdr.rows[0].cells

    p = right.paragraphs[0]
    txt(p, "דר' קגנוביץ גלינה", 10, color=BLUE, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p2 = right.add_paragraph()
    txt(p2, "משפחה, פנימית וכללית", 10, color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.RIGHT)

    pc = center.paragraphs[0]
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    txt(pc, "מכבי", 14, bold=True, color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.CENTER)
    pc2 = center.add_paragraph()
    txt(pc2, "שירותי בריאות", 14, bold=True, color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.CENTER)

    pl = left.paragraphs[0]
    txt(pl, "תאריך: 31/08/2026", 10, color=BLUE, align=WD_ALIGN_PARAGRAPH.LEFT)
    pl2 = left.add_paragraph()
    txt(pl2, "מ.ר גורם מפנה:", 9, color=BLUE, align=WD_ALIGN_PARAGRAPH.LEFT)
    pl3 = left.add_paragraph()
    txt(pl3, "024125", 11, bold=True, color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.LEFT)
    if os.path.exists(STAMP):
        stamp_row = doc.add_table(1, 3)
        no_border(stamp_row)
        sp = stamp_row.rows[0].cells[1].paragraphs[0]
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.add_run().add_picture(STAMP, width=Cm(3.2))

    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(2)

    for line in [
        "טלפון: 08-9257689",
        "פקס: 08-9257689",
        "כתובת: שטרן יאיר 14, רמלה",
    ]:
        p = doc.add_paragraph()
        txt(p, line, 10, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── Patient box ───────────────────────────────────────────────
    box = doc.add_table(4, 3)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    blue_border(box)

    hdr_row = box.rows[0]
    txt(hdr_row.cells[2].paragraphs[0], "פרטי הנבדק:", 10, bold=True, color=BLUE)
    txt(hdr_row.cells[0].paragraphs[0], "0215359910", 10, bold=True, color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.LEFT)

    patient = [
        ("שם משפחה: גורפינקל", "שם פרטי: רון", "ת.ז.: 215359910"),
        ("ת.לידה: 30/05/2005", "מין: ז", "טלפון: 08-9244787"),
        ("כתובת: הרב אפריאט 6, רמלה", "טל. עבודה/נייד: 0522460193", "מיקוד 7228212"),
    ]
    for ri, row_cells in enumerate(patient, start=1):
        for ci, val in enumerate(row_cells):
            p = box.rows[ri].cells[2 - ci].paragraphs[0]  # RTL column order
            txt(p, val, 10, color=BLUE)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ── Title + body ──────────────────────────────────────────────
    p = doc.add_paragraph()
    txt(p, "אישור מחלה", 14, bold=True, underline=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    p = doc.add_paragraph()
    txt(p, "ET4D-CV5-ZKC-6R82-449XH", 9, color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    p = doc.add_paragraph()
    txt(
        p,
        'הנני לאשר כי הנ"ל חלה/תה, האבחנה מפורטת ברשומה הרפואית.',
        10,
        color=RGBColor(0, 0, 0),
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    txt(
        p,
        'אינו/ה מסוגל/ת לעבוד מיום: 24/08/2026  עד יום: 27/08/2026  סה"כ: 4 ימים.',
        10,
        color=RGBColor(0, 0, 0),
    )

    for _ in range(3):
        doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────────
    foot = doc.add_table(2, 2)
    no_border(foot)
    txt(foot.rows[1].cells[0].paragraphs[0], "חתימה וחותמת הרופא", 10, color=BLUE)
    txt(foot.rows[0].cells[1].paragraphs[0], "31/08/2026", 10, color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.RIGHT)
    txt(foot.rows[1].cells[1].paragraphs[0], "תאריך", 10, color=BLUE, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(OUT)
    print("Saved", OUT)


if __name__ == "__main__":
    build()
